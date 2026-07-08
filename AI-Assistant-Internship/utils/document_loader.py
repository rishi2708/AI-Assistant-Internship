"""Document loading utilities for PDF, TXT, DOCX, MedQuAD, and arXiv data."""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from utils.text_processing import clean_text

logger = logging.getLogger(__name__)


def load_text_file(path: Path) -> str:
    """Read text from a UTF-8 compatible text file."""

    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    """Extract text from a PDF file."""

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def load_docx(path: Path) -> str:
    """Extract text from a DOCX file."""

    doc = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def load_uploaded_document(path: Path) -> Document:
    """Load an uploaded PDF, TXT, MD, or DOCX document."""

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = load_pdf(path)
    elif suffix in {".txt", ".md"}:
        text = load_text_file(path)
    elif suffix == ".docx":
        text = load_docx(path)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")

    return Document(
        page_content=clean_text(text),
        metadata={"source": path.name, "path": str(path), "type": suffix.lstrip(".")},
    )


def iter_json_records(path: Path) -> Iterable[dict]:
    """Yield records from a JSON, JSONL, or CSV file."""

    suffix = path.suffix.lower()
    if suffix == ".jsonl":
        with path.open("r", encoding="utf-8") as handle:
            for line in handle:
                if line.strip():
                    yield json.loads(line)
    elif suffix == ".json":
        payload = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(payload, list):
            yield from payload
        elif isinstance(payload, dict):
            records = payload.get("data") or payload.get("records") or [payload]
            yield from records
    elif suffix == ".csv":
        for row in pd.read_csv(path).fillna("").to_dict(orient="records"):
            yield row
    else:
        raise ValueError(f"Unsupported structured data file: {path}")


def load_medquad_documents(data_dir: Path) -> list[Document]:
    """Load MedQuAD-like question-answer records from a directory."""

    documents: list[Document] = []
    for path in sorted(data_dir.glob("*")):
        if path.suffix.lower() not in {".json", ".jsonl", ".csv", ".txt"}:
            continue
        if path.suffix.lower() == ".txt":
            text = clean_text(load_text_file(path))
            if text:
                documents.append(Document(page_content=text, metadata={"source": path.name}))
            continue
        for record in iter_json_records(path):
            question = record.get("question") or record.get("Question") or ""
            answer = record.get("answer") or record.get("Answer") or record.get("content") or ""
            focus = record.get("focus_area") or record.get("disease") or record.get("title") or ""
            text = clean_text(f"Question: {question}\nAnswer: {answer}")
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": path.name,
                            "question": question,
                            "focus_area": focus,
                            "dataset": "MedQuAD",
                        },
                    )
                )
    logger.info("Loaded %s MedQuAD documents", len(documents))
    return documents


def load_arxiv_documents(data_dir: Path, max_records: int | None = None) -> list[Document]:
    """Load arXiv metadata records from JSON/JSONL/CSV files."""

    documents: list[Document] = []
    for path in sorted(data_dir.glob("*")):
        if path.suffix.lower() not in {".json", ".jsonl", ".csv", ".txt"}:
            continue
        if path.suffix.lower() == ".txt":
            text = clean_text(load_text_file(path))
            if text:
                documents.append(Document(page_content=text, metadata={"source": path.name, "dataset": "arXiv"}))
            continue
        for record in iter_json_records(path):
            title = record.get("title", "")
            abstract = record.get("abstract") or record.get("summary") or ""
            authors = record.get("authors", "")
            categories = record.get("categories", "")
            arxiv_id = record.get("id") or record.get("arxiv_id") or ""
            text = clean_text(
                f"Title: {title}\nAuthors: {authors}\nCategories: {categories}\nAbstract: {abstract}"
            )
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": path.name,
                            "title": title,
                            "authors": authors,
                            "categories": categories,
                            "arxiv_id": arxiv_id,
                            "dataset": "arXiv",
                        },
                    )
                )
            if max_records and len(documents) >= max_records:
                logger.info("Loaded %s arXiv documents", len(documents))
                return documents
    logger.info("Loaded %s arXiv documents", len(documents))
    return documents
